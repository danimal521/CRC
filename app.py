from git.repo.base import Repo
import os
from docx import Document
from datetime import datetime
import openai
from pathlib import Path


def GenerateFileName():
    return str(datetime.now().strftime("%Y%m%d-%H%M%S") + ".docx")


def TryCreateReportFolder():
    print("Try to create report folder: CRCReport")
    try:
        os.mkdir("git/CRCReport", 0o666)
    # os.rmdir("git")
    except:
        print("Already there!")


def GitPush():
    try:
        repo = Repo("git")
        repo.git.add(".")
        repo.index.commit("CRCReport update")
        origin = repo.remote(name="origin")
        origin.push()
    except:
        print("Some error occured while pushing the code")


def GetTextFromInputFile(strFileName):
    with open(strFileName) as f:
        return f.read().rstrip()


def ReviewCode(strFileName):
    print("ReviewCode: " + strFileName)
    strReturn = ""
    openai.api_type = "azure"
    openai.api_version = "2023-03-15-preview"
    openai.api_base = env_openapi_base
    openai.api_key = env_openai_key

    try:
        jsonRequest = (
            {
                "role": "user",
                "content": "please do a full indepth code review of the code below:"
                + GetTextFromInputFile(strFileName),
            },
        )

        # WriteJSON(jsonRequest)

        response = openai.ChatCompletion.create(
            engine=env_openai_model,  # The deployment name you chose when you deployed the ChatGPT or GPT-4 model.
            messages=jsonRequest,
        )

        strReturn = response["choices"][0]["message"]["content"]
    except:
        print("Could not review: " + strFileName)

    return strReturn


def RateCode(strFileName):
    print("RateCode: " + strFileName)
    strReturn = ""

    openai.api_type = "azure"
    openai.api_version = "2023-03-15-preview"
    openai.api_base = env_openapi_base
    openai.api_key = env_openai_key

    try:
        jsonRequest = (
            {
                "role": "user",
                "content": "please rate the code from 1 to 10 where 10 is the worst and should be re-written, just return the number:\n\n"
                + GetTextFromInputFile(strFileName),
            },
        )

        response = openai.ChatCompletion.create(
            engine=env_openai_model,  # The deployment name you chose when you deployed the ChatGPT or GPT-4 model.
            messages=jsonRequest,
        )

        strReturn = response["choices"][0]["message"]["content"]
    except:
        print("Could not review: " + strFileName)

    return strReturn

env_openapi_base = os.environ['env_openapi_base']
env_openai_key = os.environ['env_openai_key']
env_openai_model = os.environ['env_openai_model']
env_repo = os.environ['env_repo']
env_pat = os.environ['env_pat']

HTTPS_REMOTE_URL = "https://" + env_pat + ":x-oauth-basic@" + env_repo

if os.path.exists("git"):
    print("Repo found, pulling latest")
    Repo("git").remotes.origin.pull()
else:
    print("No repo, cloning")
    Repo.clone_from(HTTPS_REMOTE_URL, "git")

TryCreateReportFolder()

print("Create report")
document = Document()
document.add_picture('28.png')

for child in Path("git").iterdir():
    if child.is_file():
        if ".cs" in Path(child).suffix:
            strFileName = "git/" + child.name
            document.add_paragraph(strFileName)
            document.add_paragraph("Need to rewrite: " + RateCode(strFileName))
            document.add_paragraph(ReviewCode(strFileName))
            document.add_page_break()

document.save("git/CRCReport/" + GenerateFileName())

print("Push repo")
GitPush()
print("Done!")


#https://python-docx.readthedocs.io/en/latest/user/quickstart.html